#!/usr/bin/env python3
"""
Mother Brain Document Generator
Generates DOCX documents from templates using dynamic ROL variable substitution.
"""
import os
import re
import json
from docx import Document

BASE_DIR = "/home/jay/.openclaw/workspace/projects/whatsapp-bot/guru-source"
TEMPLATES_DIR = os.path.join(BASE_DIR, "LATEST_UPDATES/LATEST UPDATES")
OUTPUT_DIR = os.path.join(BASE_DIR, "templates/output")

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Known ROL codes
ROL_CODES = [
    "VENDEDOR", "COMPRADOR", "ARRENDADOR", "ARRENDATARIO",
    "DONANTE", "DONATARIO", "TESTIGO", "CONYUGE",
    "APODERADO", "PODERDANTE", "GARANTE", "BENEFICIARIO",
    "DEMANDANTE", "DEMANDADO", "RECURRENTE", "RECURRIDO",
    "ACREEDOR", "DEUDOR", "SOLICITANTE", "COMPARECIENTE",
    "REPRESENTANTE", "ABOGADO", "NOTARIO", "ALGUACIL",
    "IMPUTADO", "VICTIMA", "MAGISTRADO", "REQUIRENTE",
]


def extract_required_fields(template_path):
    """Extract all {{VARIABLE}} placeholders from a template."""
    full_path = os.path.join(TEMPLATES_DIR, template_path)
    if not os.path.exists(full_path):
        return []
    doc = Document(full_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    return sorted(set(re.findall(r"\{\{([^}]+)\}\}", text)))


def analyze_template_roles(template_path):
    """
    Analyze a template and identify what ROL roles it requires.
    Returns dict: {role_code: [field1, field2, ...]}
    """
    variables = extract_required_fields(template_path)
    roles_needed = {}
    for var in variables:
        for rol in ROL_CODES:
            # Match FIELD_ROLE or FIELD_ROLE O ALIAS or FIELD_ROLE N
            patterns = [
                rf"^(.*)_{re.escape(rol)}\s+O\s+.*$",
                rf"^(.*)_{re.escape(rol)}\s+(\d+)$",
                rf"^(.*)_{re.escape(rol)}$",
            ]
            for pat in patterns:
                m = re.match(pat, var)
                if m:
                    if rol not in roles_needed:
                        roles_needed[rol] = set()
                    base = m.group(1).strip()
                    roles_needed[rol].add(base)
                    break
            else:
                continue
            break
    return {k: sorted(list(v)) for k, v in roles_needed.items()}


def match_placeholder_to_data(placeholder, collected_data, assigned_roles):
    """
    Try to match a placeholder like 'NOMBRE_VENDEDOR O PRIMERA PARTE'
    to actual data from collected_data or assigned_roles.
    Returns the value or None if no match.
    """
    # 1. Direct match in collected_data
    if placeholder in collected_data:
        return collected_data[placeholder]

    # 2. Try ROL matching
    for rol in ROL_CODES:
        # Pattern: FIELD_ROLE O ALIAS or FIELD_ROLE N or FIELD_ROLE
        m = re.match(rf"^(.*)_{re.escape(rol)}(\s+O\s+.*|\s+\d+)?$", placeholder)
        if m:
            base_field = m.group(1).strip()
            role_key = rol
            # Check if we have data for this role
            if role_key in assigned_roles:
                role_data = assigned_roles[role_key]
                # Try exact base field match
                if base_field in role_data:
                    return role_data[base_field]
                # Try with spaces instead of underscores
                base_field_spaced = base_field.replace("_", " ")
                if base_field_spaced in role_data:
                    return role_data[base_field_spaced]
            break

    # 3. Try with spaces instead of underscores (for direct vars)
    placeholder_spaced = placeholder.replace("_", " ")
    if placeholder_spaced in collected_data:
        return collected_data[placeholder_spaced]

    return None


def replace_in_paragraph(paragraph, replacements):
    """Replace all {{VARIABLE}} placeholders in a paragraph."""
    if not paragraph.runs:
        return
    full_text = "".join([run.text for run in paragraph.runs])
    found = re.findall(r"\{\{[^}]+\}\}", full_text)
    if not found:
        return

    new_text = full_text
    for placeholder in found:
        if placeholder in replacements:
            new_text = new_text.replace(placeholder, replacements[placeholder])

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_in_table(table, replacements):
    """Replace placeholders in all table cells."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)


def generate_document(template_path, collected_data, assigned_roles, output_name=None):
    """
    Generate a document from a template.
    """
    full_path = os.path.join(TEMPLATES_DIR, template_path)
    if not os.path.exists(full_path):
        raise FileNotFoundError(f"Template not found: {full_path}")

    doc = Document(full_path)

    # Extract all placeholders from the template
    all_placeholders = set()
    for paragraph in doc.paragraphs:
        all_placeholders.update(re.findall(r"\{\{[^}]+\}\}", " ".join([r.text for r in paragraph.runs])))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    all_placeholders.update(re.findall(r"\{\{[^}]+\}\}", " ".join([r.text for r in paragraph.runs])))

    # Build replacements by matching each actual placeholder
    replacements = {}
    for ph in all_placeholders:
        inner = ph[2:-2]  # Remove {{ and }}
        value = match_placeholder_to_data(inner, collected_data, assigned_roles)
        if value is not None:
            replacements[ph] = str(value)

    # Apply replacements
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        replace_in_table(table, replacements)

    # Save output
    if not output_name:
        base = os.path.splitext(os.path.basename(template_path))[0]
        output_name = f"{base}_GENERADO.docx"

    output_path = os.path.join(OUTPUT_DIR, output_name)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    test_path = "CONTRATOS CIVILES/CONTRATOS BAJO FIRMA PRIVADA/BIENES TRASLATIVOS/VEHICULOS/ACTO DE VENTA DE VEHICULO.docx"
    print("Testing generation...")
    print("Required fields:", extract_required_fields(test_path))
    print("Roles:", analyze_template_roles(test_path))

    result = generate_document(
        test_path,
        collected_data={
            "CIUDAD_FIRMA": "Santo Domingo",
            "DIA_TEXTO": "quince",
            "DIA_NUMERO": "15",
            "MES_TEXTO": "mayo",
            "AÑO_TEXTO": "dos mil veintiséis",
            "AÑO_NUMERO": "2026",
            "DESCRIPCION_DEL_BIEN": "Toyota Corolla 2020, placa A123456, color blanco",
            "PRECIO_VENTA_LETRAS": "Doscientos cincuenta mil pesos",
            "PRECIO_VENTA_NUMEROS": "250,000.00",
            "NUMERO_MATRICULA": "123456789",
            "FECHA_EXPEDICION_MATRICULA": "10 de enero de 2020",
        },
        assigned_roles={
            "VENDEDOR": {
                "NOMBRE": "JUAN PEREZ GARCIA",
                "NACIONALIDAD": "Dominicano",
                "DOCUMENTO IDENTIDAD": "001-1234567-8",
                "DIRECCION O DOMICILIO": "Calle Principal #123, Santo Domingo",
            },
            "COMPRADOR": {
                "NOMBRE": "MARIA LOPEZ SANTOS",
                "NACIONALIDAD": "Dominicana",
                "DOCUMENTO IDENTIDAD": "001-8765432-1",
                "DIRECCION O DOMICILIO": "Avenida Independencia #456, Santo Domingo",
            },
        },
    )
    print(f"Generated: {result}")
